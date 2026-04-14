"""
syllabusParser.py
=================
Parses a Syllabus DOCX and/or a Micro Syllabus DOCX and returns a unified
JSON structure:

{
  "units": [
    {
      "unitNumber": 1,
      "topics": [
        {
          "title": "Introduction to Software Engineering",
          "subtopics": ["...", "..."]
        }
      ]
    }
  ]
}

Usage:
    python syllabusParser.py --syllabus <path> [--micro <path>]

Rules:
- All output (JSON) goes to stdout.
- All warnings/debug go to stderr so they don't corrupt JSON.
- Only python-docx and stdlib — no AI/ML libraries.
"""

import sys
import json
import re
import argparse
import unicodedata

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print(json.dumps({"error": "python-docx is not installed. Run: pip install python-docx"}), file=sys.stdout)
    sys.exit(1)


# ── Roman numeral → int map ──────────────────────────────────────────────────
ROMAN_MAP = {
    'I': 1, 'II': 2, 'III': 3, 'IV': 4, 'V': 5,
    'VI': 6, 'VII': 7, 'VIII': 8, 'IX': 9, 'X': 10,
}


def roman_to_int(s):
    """Convert Roman numeral string to int, or return None if not a valid Roman numeral."""
    return ROMAN_MAP.get(s.strip().upper())


def clean_co_tags(text):
    """Remove CO-mapping text like CO1, CO2, CO1,CO2 from the end (or anywhere) in a string."""
    # Remove patterns like "CO1", "CO 2", "CO1,CO2", "(CO1)", etc.
    text = re.sub(r'\(?CO\s*\d+(?:\s*,\s*CO\s*\d+)*\)?', '', text, flags=re.IGNORECASE)
    # Clean up leftover trailing punctuation / whitespace
    text = re.sub(r'[\s,;:]+$', '', text)
    return text.strip()


def normalize_title(title):
    """Lowercase, strip whitespace, strip trailing colon — for fuzzy matching."""
    t = title.lower().strip()
    t = t.rstrip(':').strip()
    # Remove extra whitespace
    t = re.sub(r'\s+', ' ', t)
    return t


def is_bold_paragraph(para):
    """
    A paragraph is considered BOLD if:
    - Its style name explicitly contains 'Heading', OR
    - ALL non-empty runs in the paragraph have run.bold == True, OR
    - ANY run is bold AND the paragraph text ends with ':'
    """
    style_name = (para.style.name or '').lower()
    if 'heading' in style_name:
        return True

    runs = [r for r in para.runs if r.text.strip()]
    if not runs:
        return False

    # All runs bold
    if all(r.bold for r in runs):
        return True

    # Any run bold + ends with colon (common pattern in syllabi)
    any_bold = any(r.bold for r in runs)
    ends_colon = para.text.strip().endswith(':')
    if any_bold and ends_colon:
        return True

    return False


def get_cell_paragraphs(cell):
    """Return all paragraphs from a table cell, including nested tables if any."""
    return cell.paragraphs





def split_subtopic_items(text):
    """
    Split a subtopic string on ', ' (comma + space) but ONLY when:
      1. Not inside parentheses (paren-depth == 0), AND
      2. The current segment has NOT yet seen a colon.

    Once a colon appears in the current segment, any further commas are treated
    as part of that segment's inline sub-enumeration and are NOT split points.

    Examples
    --------
    'Waterfall model, Incremental process model, Evolutionary process models: prototyping, spiral'
    → ['Waterfall model', 'Incremental process model',
       'Evolutionary process models: prototyping, spiral']

    'Software Process Framework (The process framework, umbrella activities)'
    → ['Software Process Framework (The process framework, umbrella activities)']

    'Agile Process (XP, Scrum), Spiral model, RAD model'
    → ['Agile Process (XP, Scrum)', 'Spiral model', 'RAD model']
    """
    items = []
    depth = 0       # paren nesting depth
    has_colon = False   # has the current segment seen a ':' at depth-0?
    current = []

    i = 0
    while i < len(text):
        ch = text[i]
        if ch in '([{':
            depth += 1
            current.append(ch)
        elif ch in ')]}':
            depth = max(0, depth - 1)
            current.append(ch)
        elif ch == ':' and depth == 0:
            has_colon = True
            current.append(ch)
        elif ch == ',' and depth == 0 and not has_colon:
            # Split only at top-level commas in segments that haven't seen a colon
            if i + 1 < len(text) and text[i + 1] == ' ':
                item = ''.join(current).strip()
                if item:
                    items.append(item)
                current = []
                has_colon = False
                i += 2   # skip ', '
                continue
            else:
                current.append(ch)
        else:
            current.append(ch)
        i += 1

    # Flush the last segment
    tail = ''.join(current).strip()
    if tail:
        items.append(tail)

    return items if items else [text]


def parse_cell_content(cell, paragraphs=None):
    """
    Parse paragraphs from a table cell (column 1 of a unit row).
    Accepts an optional `paragraphs` list so continuation-row paragraphs can
    be appended before calling this function.
    Returns a list of { 'title': str, 'subtopics': [str] }.
    """
    topics = []
    current_title = None
    current_subtopics = []

    para_list = paragraphs if paragraphs is not None else cell.paragraphs

    for para in para_list:
        raw_text = para.text
        if not raw_text:
            continue

        text = clean_co_tags(raw_text).strip()
        if not text:
            continue

        bold = is_bold_paragraph(para)
        ends_colon = text.endswith(':')

        # It is a topic title if it's bold OR ends with ':'
        if bold or ends_colon:
            # Save the previous group
            if current_title is not None:
                topics.append({
                    'title': current_title.rstrip(':').strip(),
                    'subtopics': current_subtopics,
                })
            current_title = text.rstrip(':').strip()
            current_subtopics = []
        else:
            # It's a sub-topic bullet — split on ', ' outside parens
            sub_items = split_subtopic_items(text)
            if current_title is not None:
                current_subtopics.extend(sub_items)
            else:
                # No title seen yet — treat first non-bold paragraph as a fallback title
                current_title = text
                current_subtopics = []

    # Flush last group
    if current_title is not None:
        topics.append({
            'title': current_title.rstrip(':').strip(),
            'subtopics': current_subtopics,
        })

    return topics


# ── SYLLABUS PARSER ──────────────────────────────────────────────────────────

def find_syllabus_table(doc):
    """
    Find the table in the syllabus DOCX that has 'Unit No.' as a header cell.
    Returns the table object or None.
    """
    for table in doc.tables:
        if not table.rows:
            continue
        first_row = table.rows[0]
        for cell in first_row.cells:
            cell_text = cell.text.strip()
            if cell_text == 'Unit No.' or cell_text.lower() == 'unit no.':
                return table
    return None


def parse_syllabus(docx_path):
    """
    Parse the Syllabus DOCX.
    Returns dict: { unit_number (int): [{ title, subtopics }] }
    """
    print(f"[syllabusParser] Parsing syllabus: {docx_path}", file=sys.stderr)
    doc = Document(docx_path)

    table = find_syllabus_table(doc)
    if table is None:
        print("[syllabusParser] WARNING: Could not find the 'Unit No.' table in syllabus.", file=sys.stderr)
        return {}

    units = {}

    # Skip the header row (row 0)
    for row in table.rows[1:]:
        cells = row.cells
        if len(cells) < 2:
            continue

        unit_col_text = cells[0].text.strip()
        unit_num = roman_to_int(unit_col_text)

        if unit_num is None:
            # Not a unit row — skip
            print(f"[syllabusParser] Skipping syllabus row with col0='{unit_col_text}'", file=sys.stderr)
            continue

        content_cell = cells[1]
        topics = parse_cell_content(content_cell)

        if not topics:
            print(f"[syllabusParser] Unit {unit_num}: no topics parsed from syllabus cell.", file=sys.stderr)
            continue

        units[unit_num] = topics

    return units


# ── MICRO SYLLABUS PARSER ────────────────────────────────────────────────────

UNIT_ROW_RE = re.compile(r'^UNIT[-\s]*(\d+)', re.IGNORECASE)

# Rows whose col0 matches these are noise — stop absorption and skip entirely
_NOISE_KEYWORDS = re.compile(
    r'^(learning\s+resource|text\s*book|course\s+co.?ordinator|hod|'
    r'e.?resource|reference|web\s+link|co\s+mapping|'
    r'module\s+coordinator|program\s+coordinator|co\s+\d)',
    re.IGNORECASE
)

# CO-row detector: col0 like "CO1, CO2" or "CO1,CO2,CO4"
_CO_ROW_RE = re.compile(r'^CO\s*\d+', re.IGNORECASE)


class _FakeStyle:
    """Minimal style stand-in — has no heading marker."""
    name = ''


class _FakeRun:
    """Minimal run stand-in carrying text + bold flag."""
    def __init__(self, text, bold=False):
        self.text = text
        self.bold = bold


class _FakeParagraph:
    """
    Minimal paragraph stand-in so that parse_cell_content can process
    section sub-headings that live in col0 of a continuation row
    (e.g. "PROCESS MODELS:", "TESTING CONVENTIONAL APPLICATIONS:").
    """
    def __init__(self, text, bold=True):
        self.text = text
        self.style = _FakeStyle()
        self.runs = [_FakeRun(text, bold=bold)]


def is_micro_unit_row(cell_text):
    """Return unit number (int) if col 0 of a micro syllabus row is a UNIT-N header, else None."""
    m = UNIT_ROW_RE.match(cell_text.strip())
    if m:
        return int(m.group(1))
    return None


def parse_micro_syllabus(docx_path):
    """
    Parse the Micro Syllabus DOCX.
    Returns dict: { unit_number (int): [{ title, subtopics }] }

    Edge-cases handled
    ------------------
    1. Split unit rows   — a single unit's content may be spread across several
       rows where col0 of later rows is blank.  We absorb those rows.
    2. Sub-section heads — col0 of a continuation row may contain a section
       subheading like "PROCESS MODELS:" or "TESTING CONVENTIONAL APPLICATIONS:".
       These are NOT unit boundaries; we inject a fake bold paragraph so they
       become topic titles, then absorb col1 as their subtopics.
    3. CO-mapping rows  — col0 rows that start with "CO1, CO2" etc. are skipped
       silently (they carry no curriculum content).
    4. Noise rows       — Learning Resources, Text Books, HOD etc. stop absorption.
    5. Multiple tables  — scans all top-level tables to find the one with UNIT-N.
    """
    print(f"[syllabusParser] Parsing micro syllabus: {docx_path}", file=sys.stderr)
    doc = Document(docx_path)

    if not doc.tables:
        print("[syllabusParser] WARNING: No tables found in micro syllabus.", file=sys.stderr)
        return {}

    print(f"[syllabusParser] Document has {len(doc.tables)} table(s).", file=sys.stderr)

    # Flatten all rows from all tables into a single list.
    # This handles the edge case where the author split the units across multiple separate tables 
    # (e.g. by inserting a paragraph or page break between units).
    rows = []
    for table in doc.tables:
        rows.extend(table.rows)

    print(f"[syllabusParser] Flattened {len(rows)} rows across all tables.", file=sys.stderr)
    units = {}

    i = 0
    while i < len(rows):
        row = rows[i]
        cells = row.cells

        if len(cells) < 2:
            i += 1
            continue

        col0_text = cells[0].text.strip()
        unit_num = is_micro_unit_row(col0_text)

        if unit_num is None:
            # Not a unit start row — skip header, CO outcomes table, etc.
            i += 1
            continue

        print(f"[syllabusParser] Found UNIT-{unit_num} at row {i}.", file=sys.stderr)

        # Collect all paragraphs for this unit starting with the current row's col1
        all_paragraphs = list(cells[1].paragraphs)
        i += 1

        # ── Absorb continuation rows ─────────────────────────────────────────
        while i < len(rows):
            cont_cells = rows[i].cells
            if len(cont_cells) < 2:
                i += 1
                continue

            cont_col0 = cont_cells[0].text.strip()

            # STOP: next UNIT-N starts
            if is_micro_unit_row(cont_col0) is not None:
                print(f"[syllabusParser] Unit {unit_num} ends before row {i} (next unit).", file=sys.stderr)
                break

            # STOP: explicit noise row (Learning Resources, HOD, etc.)
            if cont_col0 and _NOISE_KEYWORDS.match(cont_col0):
                print(f"[syllabusParser] Unit {unit_num} ends before row {i} — noise: '{cont_col0[:40]}'.", file=sys.stderr)
                break

            # SKIP silently: CO-mapping rows like "CO1, CO2, CO4"
            if cont_col0 and _CO_ROW_RE.match(cont_col0):
                print(f"[syllabusParser] Skipping CO row: '{cont_col0[:40]}'.", file=sys.stderr)
                i += 1
                continue

            # ABSORB: col0 is blank — straight continuation
            if not cont_col0:
                all_paragraphs.extend(cont_cells[1].paragraphs)
                i += 1
                continue

            # ABSORB: col0 has a section sub-heading like "PROCESS MODELS:" or
            # "TESTING CONVENTIONAL APPLICATIONS:" — these are topic groups
            # within this unit.  Inject as a bold fake paragraph so it becomes
            # a topic title, then absorb col1 content as its subtopics.
            heading_text = cont_col0.rstrip(':').strip()
            print(f"[syllabusParser] Unit {unit_num}: absorbing sub-section '{heading_text}' from col0 at row {i}.", file=sys.stderr)
            all_paragraphs.append(_FakeParagraph(heading_text + ':', bold=True))
            all_paragraphs.extend(cont_cells[1].paragraphs)
            i += 1
            continue
        # ── End absorption ───────────────────────────────────────────────────

        # Parse the full consolidated paragraph list for this unit
        topics = parse_cell_content(None, paragraphs=all_paragraphs)

        if not topics:
            print(f"[syllabusParser] Unit {unit_num}: no topics parsed.", file=sys.stderr)
            continue

        print(f"[syllabusParser] Unit {unit_num}: {len(topics)} topic(s) parsed.", file=sys.stderr)
        units[unit_num] = topics

    return units


# ── MERGE ────────────────────────────────────────────────────────────────────

def fuzzy_match(title_a, title_b):
    """
    Simple token-overlap fuzzy match.
    Returns True if the two titles are 'similar enough'.
    Both titles should already be normalize_title()'d.
    """
    if title_a == title_b:
        return True

    # One is a prefix of the other (handles colon trimming differences)
    if title_a.startswith(title_b) or title_b.startswith(title_a):
        return True

    # Token overlap — if 70%+ tokens match
    tokens_a = set(title_a.split())
    tokens_b = set(title_b.split())
    if not tokens_a or not tokens_b:
        return False

    intersection = tokens_a & tokens_b
    union = tokens_a | tokens_b
    if len(union) == 0:
        return False

    jaccard = len(intersection) / len(union)
    return jaccard >= 0.6


def merge_units(syllabus_units, micro_units):
    """
    Merge syllabus and micro syllabus structures.
    Syllabus = source of truth for topic titles & unit list.
    Micro syllabus = source of sub-topics.
    Returns list of { unitNumber, topics: [{ title, subtopics }] }
    """
    result = []

    all_unit_numbers = sorted(set(list(syllabus_units.keys()) + list(micro_units.keys())))

    for unit_num in all_unit_numbers:
        syl_topics = syllabus_units.get(unit_num, [])
        micro_topics = micro_units.get(unit_num, [])

        # Build a lookup map from normalized micro title → subtopics
        micro_lookup = {}
        for mt in micro_topics:
            norm = normalize_title(mt['title'])
            micro_lookup[norm] = mt['subtopics']

        merged_topics = []

        if syl_topics:
            # Use syllabus as source of truth for titles
            for st in syl_topics:
                norm_st = normalize_title(st['title'])

                # Try to find a matching micro topic
                matched_subtopics = st.get('subtopics', [])  # Fallback: subtopics from syllabus itself

                for norm_mt, subtopics in micro_lookup.items():
                    if fuzzy_match(norm_st, norm_mt):
                        if subtopics:
                            matched_subtopics = subtopics
                        break

                merged_topics.append({
                    'title': st['title'],
                    'subtopics': matched_subtopics,
                })
        else:
            # No syllabus for this unit — use micro directly
            for mt in micro_topics:
                merged_topics.append({
                    'title': mt['title'],
                    'subtopics': mt['subtopics'],
                })

        if merged_topics:
            result.append({
                'unitNumber': unit_num,
                'topics': merged_topics,
            })

    return result


def parse_single(units_dict):
    """Convert a single-source units dict to the output list format."""
    result = []
    for unit_num in sorted(units_dict.keys()):
        topics = units_dict[unit_num]
        if topics:
            result.append({
                'unitNumber': unit_num,
                'topics': topics,
            })
    return result


# ── MAIN ─────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description='Parse Syllabus and/or Micro Syllabus DOCX files.')
    parser.add_argument('--syllabus', type=str, default=None, help='Path to Syllabus DOCX file')
    parser.add_argument('--micro', type=str, default=None, help='Path to Micro Syllabus DOCX file')
    args = parser.parse_args()

    if not args.syllabus and not args.micro:
        print(json.dumps({'error': 'At least one of --syllabus or --micro must be provided.'}))
        sys.exit(1)

    syllabus_units = {}
    micro_units = {}

    if args.syllabus:
        try:
            syllabus_units = parse_syllabus(args.syllabus)
        except Exception as e:
            print(f"[syllabusParser] ERROR parsing syllabus: {e}", file=sys.stderr)

    if args.micro:
        try:
            micro_units = parse_micro_syllabus(args.micro)
        except Exception as e:
            print(f"[syllabusParser] ERROR parsing micro syllabus: {e}", file=sys.stderr)

    if syllabus_units and micro_units:
        units_list = merge_units(syllabus_units, micro_units)
    elif syllabus_units:
        units_list = parse_single(syllabus_units)
    elif micro_units:
        units_list = parse_single(micro_units)
    else:
        print(json.dumps({'error': 'No units could be extracted from the provided files.'}))
        sys.exit(1)

    output = {'units': units_list}
    print(json.dumps(output, ensure_ascii=False, indent=2))


if __name__ == '__main__':
    main()
